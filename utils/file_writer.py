"""
多格式输出器 — 支持 JSONL / TXT / Word (.docx) 输出。
"""

from __future__ import annotations

import json
import logging
from pathlib import Path

from .schema import OutputItem

logger = logging.getLogger(__name__)


def detect_output_format(path: Path) -> str:
    suffix = path.suffix.lower()
    format_map = {
        ".jsonl": "jsonl",
        ".json": "jsonl",
        ".txt": "txt",
        ".docx": "docx",
        ".doc": "docx",
    }
    fmt = format_map.get(suffix)
    if not fmt:
        raise ValueError(f"不支持的输出格式: {suffix}（支持 .jsonl / .txt / .docx）")
    return fmt


def write_output(items: list[OutputItem], path: Path) -> None:
    """根据文件后缀自动选择输出格式。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    fmt = detect_output_format(path)

    if fmt == "jsonl":
        _write_jsonl(items, path)
    elif fmt == "txt":
        _write_txt(items, path)
    elif fmt == "docx":
        _write_docx(items, path)

    logger.info("已写入 %s (%d 条, 格式: %s)", path, len(items), fmt)


def _write_jsonl(items: list[OutputItem], path: Path) -> None:
    with open(path, "w", encoding="utf-8") as f:
        for item in items:
            f.write(item.model_dump_json(ensure_ascii=False) + "\n")


def _write_txt(items: list[OutputItem], path: Path) -> None:
    with open(path, "w", encoding="utf-8") as f:
        for i, item in enumerate(items):
            f.write(f"{'='*60}\n")
            f.write(f"[{item.id}]\n\n")
            f.write(f"【问题】\n{item.question}\n\n")
            f.write(f"【回答】\n{item.answer}\n")

            if item.evaluation:
                ev = item.evaluation
                f.write(f"\n【评分】{ev.total_score:.1f}/10 ({'通过' if ev.passed else '未通过'})\n")
                for d in ev.dimensions:
                    f.write(f"  - {d.dimension}: {d.score:.1f} — {d.reason}\n")
                if ev.suggestion and ev.suggestion != "无":
                    f.write(f"【建议】{ev.suggestion}\n")

            if item.rework_count > 0:
                f.write(f"【重写次数】{item.rework_count}\n")

            f.write("\n")


def _write_docx(items: list[OutputItem], path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "输出 .docx 文件需要 python-docx 库，请运行: pip install python-docx"
        )

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    title = doc.add_heading('语料处理结果', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"共 {len(items)} 条数据")
    doc.add_paragraph("")

    for item in items:
        doc.add_heading(f"[{item.id}]", level=2)

        q_para = doc.add_paragraph()
        q_run = q_para.add_run("【问题】")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph(item.question)

        a_para = doc.add_paragraph()
        a_run = a_para.add_run("【回答】")
        a_run.bold = True
        a_run.font.color.rgb = RGBColor(0, 153, 0)
        doc.add_paragraph(item.answer)

        if item.evaluation:
            ev = item.evaluation
            score_para = doc.add_paragraph()
            score_run = score_para.add_run(
                f"【评分】{ev.total_score:.1f}/10 ({'通过' if ev.passed else '未通过'})"
            )
            score_run.bold = True
            color = RGBColor(0, 153, 0) if ev.passed else RGBColor(204, 0, 0)
            score_run.font.color.rgb = color

            for d in ev.dimensions:
                doc.add_paragraph(f"  {d.dimension}: {d.score:.1f} — {d.reason}", style='List Bullet')

            if ev.suggestion and ev.suggestion != "无":
                s_para = doc.add_paragraph()
                s_run = s_para.add_run(f"【建议】{ev.suggestion}")
                s_run.italic = True

        if item.rework_count > 0:
            doc.add_paragraph(f"重写次数: {item.rework_count}")

        doc.add_page_break()

    doc.save(str(path))
