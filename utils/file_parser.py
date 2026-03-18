"""
多格式文件解析器 — 支持 JSONL / TXT / Word (.docx) 输入。

对于非结构化输入（Word/TXT），使用 LLM 智能提取 Question-Answer 对。
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

from .schema import CorpusItem

logger = logging.getLogger(__name__)


def detect_format(path: Path) -> str:
    suffix = path.suffix.lower()
    format_map = {
        ".jsonl": "jsonl",
        ".json": "jsonl",
        ".txt": "txt",
        ".docx": "docx",
        ".doc": "docx",
    }
    fmt = format_map.get(suffix)
    if not fmt:
        raise ValueError(f"不支持的文件格式: {suffix}（支持 .jsonl / .txt / .docx）")
    return fmt


def load_file(path: Path) -> list[CorpusItem] | str:
    """
    加载文件。
    - JSONL → 直接解析为 CorpusItem 列表
    - TXT/Word → 返回原始文本（需要后续 LLM 提取）
    """
    fmt = detect_format(path)

    if fmt == "jsonl":
        return _parse_jsonl(path)
    elif fmt == "txt":
        return _read_txt(path)
    elif fmt == "docx":
        return _read_docx(path)
    else:
        raise ValueError(f"不支持的格式: {fmt}")


def _parse_jsonl(path: Path) -> list[CorpusItem]:
    items: list[CorpusItem] = []
    with open(path, "r", encoding="utf-8") as f:
        for idx, line in enumerate(f):
            line = line.strip()
            if not line:
                continue
            data = json.loads(line)
            if "id" not in data:
                data["id"] = idx
            items.append(CorpusItem(**data))
    return items


def _read_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "读取 .docx 文件需要 python-docx 库，请运行: pip install python-docx"
        )

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def build_extraction_prompt(raw_text: str, max_chars: int = 15000) -> str:
    """构建让 LLM 从原始文本中提取 QA 对的 prompt。"""
    truncated = raw_text[:max_chars]
    if len(raw_text) > max_chars:
        truncated += "\n... (文本已截断)"

    return f"""\
请从以下原始文本中识别并提取所有的「问题-答案」对。

# 提取规则
1. 识别文本中所有明确或隐含的问答对
2. 如果文本有明显的 Q/A、问/答、题目/解答等标记，按标记拆分
3. 如果是表格形式（问题列 + 答案列），逐行提取
4. 如果没有明显标记，根据语义判断哪些是问题、哪些是对应的答案
5. 为每条提取结果分配一个递增 ID（从 1 开始）

# 输出格式
严格输出 JSON 数组（不要包含 ```json 标记）：
[
  {{"id": "1", "question": "提取到的问题", "answer": "提取到的答案"}},
  {{"id": "2", "question": "...", "answer": "..."}}
]

如果某条只有问题没有答案，answer 填空字符串 ""。

# 原始文本
{truncated}
"""


def parse_extracted_json(llm_output: str) -> list[CorpusItem]:
    """解析 LLM 提取结果，容错处理。"""
    text = llm_output.strip()

    json_match = re.search(r'\[.*\]', text, re.DOTALL)
    if json_match:
        text = json_match.group()

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        text = re.sub(r'```json\s*', '', text)
        text = re.sub(r'```\s*', '', text)
        data = json.loads(text)

    items = []
    for idx, entry in enumerate(data):
        if isinstance(entry, dict) and "question" in entry:
            items.append(CorpusItem(
                id=entry.get("id", str(idx + 1)),
                question=entry["question"],
                answer=entry.get("answer", ""),
            ))

    return items
